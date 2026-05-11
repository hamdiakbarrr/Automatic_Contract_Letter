import streamlit as st
import pandas as pd
from docx import Document
import os
from docx2pdf import convert
from PyPDF2 import PdfMerger
import tempfile
import shutil
import subprocess

from docx.shared import Pt # Tambahkan import ini di bagian atas script

def replace_text_logic(doc, data_map):
    import re
    from docx.shared import Pt
    
    def process_paragraphs(paragraphs):
        for p in paragraphs:
            if any(key in p.text for key in data_map):
                # 1. Ambil referensi ukuran font asli (jika ada)
                # Jika tidak terdeteksi, kita set manual ke 12pt (standar Times New Roman)
                ref_font_size = p.runs[0].font.size if p.runs and p.runs[0].font.size else Pt(9)
                
                full_text = p.text
                val_no_surat = str(data_map.get("{{no_surat}}", ""))
                
                # Ganti data dari Excel
                for key, value in data_map.items():
                    if key in full_text:
                        full_text = full_text.replace(key, str(value) if pd.notna(value) else "")

                # 2. Kosongkan runs lama
                for run in p.runs:
                    run.text = ""

                # 3. Pattern untuk Bold (PIHAK, PASAL, Nomor Surat)
                escaped_no_surat = re.escape(val_no_surat) if val_no_surat else "XYZ123_TIDAK_ADA"
                pattern = f"(PIHAK PERTAMA|PIHAK KEDUA|PASAL \d+|{escaped_no_surat})"
                
                parts = re.split(pattern, full_text)
                
                for part in parts:
                    if not part: continue
                    
                    new_run = p.add_run(part)
                    
                    # --- PENGATURAN FONT TIMES NEW ROMAN ---
                    new_run.font.name = 'Times New Roman'
                    new_run.font.size = ref_font_size
                    
                    # Logika Bold/Underline
                    if part in ["PIHAK PERTAMA", "PIHAK KEDUA"] or part.startswith("PASAL"):
                        new_run.bold = True
                    elif val_no_surat != "" and part == val_no_surat:
                        new_run.bold = True
                        new_run.underline = True
                    else:
                        new_run.bold = False

    # Jalankan proses
    process_paragraphs(doc.paragraphs)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                process_paragraphs(cell.paragraphs)

# UI Dashboard
st.set_page_config(page_title="Generator PDF Mandiri Jaya", layout="wide")
st.title("📄 Generator Kontrak Tunggal (PDF)")
st.write("Sistem akan menggabungkan semua data Excel menjadi satu file PDF.")

uploaded_template = st.file_uploader("1. Unggah Template Word (.docx)", type=["docx"])
uploaded_excel = st.file_uploader("2. Unggah Data Excel (.xlsx)", type=["xlsx"])

if st.button("Generate Satu File PDF"):
    if uploaded_template and uploaded_excel:
        try:
            # Membaca data excel
            df = pd.read_excel(uploaded_excel).dropna(how='all', axis=0)
            
            # Membuat direktori sementara
            temp_dir = tempfile.mkdtemp()
            docx_folder = os.path.join(temp_dir, "docx")
            pdf_folder = os.path.join(temp_dir, "pdf")
            os.makedirs(docx_folder)
            os.makedirs(pdf_folder)

            merger = PdfMerger()
            progress_bar = st.progress(0)
            status_text = st.empty()
            
            # Simpan file template sementara ke disk
            temp_template_path = os.path.join(temp_dir, "template.docx")
            with open(temp_template_path, "wb") as f:
                f.write(uploaded_template.getbuffer())

            for index, row in df.iterrows():
                # Buka kembali template asli untuk setiap baris data
                doc = Document(temp_template_path)
                
                # Mapping data sesuai kolom Excel
                data_map = {
                    "{{no_surat}}": row['no_surat'],
                    "{{tanggalsurat}}": row['tanggalsurat'],
                    "{{nama_p1}}": row['nama_p1'],
                    "{{jabatan_p1}}": row['jabatan_p1'],
                    "{{alamat_p1}}": row['alamat_p1'],
                    "{{nama_p2}}": row['nama_p2'],
                    "{{ttl_p2}}": row['ttl_p2'],
                    "{{jenis_kelamin_p2}}": row['jenis_kelamin_p2'],
                    "{{agama_p2}}": row['agama_p2'],
                    "{{alamat_p2}}": row['alamat_p2'],
                    "{{nik_p2}}": row['nik_p2'],
                    "{{notelfon_p2}}": row['notelfon_p2'],
                    "{{alamatkerja_p2}}": row['alamatkerja_p2'],
                    "{{tanggal_mulai_p2}}": row['tanggal_mulai_p2'],
                    "{{tanggal_berakhir_p2}}": row['tanggal_berakhir_p2'],
                    "{{alamat_p2}}": row['alamat_p2'],
                    "{{jabatan_p2}}": row['jabatan_p2'],
                    "{{nama_p3}}": row['nama_p3'],
                    "{{nama_p4}}": row['nama_p4']
                }

                # Ganti teks dengan format yang sudah diperbaiki
                replace_text_logic(doc, data_map)
                
                # Simpan DOCX hasil penggantian
                docx_path = os.path.join(docx_folder, f"temp_{index}.docx")
                doc.save(docx_path)
                
                # Konversi DOCX ke PDF
                pdf_path = os.path.join(pdf_folder, f"temp_{index}.pdf")
                status_text.text(f"Memproses ({index+1}/{len(df)}): {row['nama_p2']}")
                
                # Menjalankan konversi (Membutuhkan MS Word di Windows)
                subprocess.run(['lowriter', '--headless', '--convert-to', 'pdf', '--outdir', pdf_folder, docx_path])
                
                # Masukkan ke dalam daftar penggabungan
                merger.append(pdf_path)
                
                progress_bar.progress((index + 1) / len(df))

            # Gabungkan semua PDF menjadi satu file
            output_pdf_path = os.path.join(temp_dir, "Kontrak_Gabungan_Mandiri_Jaya.pdf")
            merger.write(output_pdf_path)
            merger.close()

            # Tampilkan tombol download
            with open(output_pdf_path, "rb") as f:
                st.success("✅ Selesai! Semua kontrak telah digabungkan menjadi satu PDF.")
                st.download_button(
                    label="📥 Download PDF Gabungan",
                    data=f,
                    file_name="Semua_Kontrak_Mandiri_Jaya.pdf",
                    mime="application/pdf"
                )
            
            # Bersihkan file sampah di folder sementara
            shutil.rmtree(temp_dir)

        except Exception as e:
            st.error(f"Terjadi kesalahan teknis: {e}")
    else:
        st.warning("Mohon unggah file Template Word dan Data Excel terlebih dahulu.")
